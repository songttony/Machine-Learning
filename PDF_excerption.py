#PDF excerpt tool
#initialize the environment
import os
import sys
import string
import importlib

#GUI interface class
import tkinter
import tkinter.messagebox
from tkinter import *

#import word document
#pip3 install python-docx to install (python-docx is compatible to Python3.x)
from docx import Document

def main():
    results = '/tmp/PDFresults.docx'
    outputs = '/tmp/PDFexcerption.docx'
    PDFfile_loc = PDFfile_input('PDF file location:')
    #parse PDF to word document
    parsePDF(PDFfile_loc, results)
    #Excerpt WORD with specified pattern, e.g. RL. RC. for this example
    ExcerptPDF(results, outputs)

#Excerpt WORD with specified pattern.
def ExcerptPDF(para_results,para_outputs):
    word_doc = Document(para_results)

    #create an empty WORD document
    outputs_doc = Document()

    #count the # of rules
    RL_counter = 0
    RC_counter = 0

    for paragraph in word_doc.paragraphs:
        if 'RL.' in paragraph.text:
            RL_counter += 1
            outputs_doc.add_paragraph(paragraph.text)
        else:
            if 'RC.' in paragraph.text:
                RC_counter += 1
                outputs_doc.add_paragraph(paragraph.text)

    outputs_doc.add_paragraph('RL Count:' + str(RL_counter))
    outputs_doc.add_paragraph('RC Count:' + str(RC_counter))
    outputs_doc.save(para_outputs)


    #code to store into TXT format
    counter = [0,0,0,0,0,0]
    RL_textstringShall = ''
    RL_textstringMust = ''
    RL_textstringOther = ''
    RC_textstringShould = ''
    RC_textstringShouldNot = ''
    RC_textstringOther = ''
    for paragraph in word_doc.paragraphs:
        if 'RL.' in paragraph.text:
            if 'shall be ' in paragraph.text:
                # counting RL-> 'shall be' items
                counter[0] += 1
                RL_textstringShall = RL_textstringShall + '\n' + paragraph.text
            else:
                if 'must not be' in paragraph.text:
                    # counting RL -> 'must not be downrated' items
                    counter[1] += 1
                    RL_textstringMust = RL_textstringMust + '\n' + paragraph.text
                else:
                    # counting RL -> other items
                    counter[2] += 1
                    RL_textstringOther = RL_textstringOther + '\n' + paragraph.text
        if 'RC.' in paragraph.text:
            if 'should be' in paragraph.text:
                # counting RC -> 'should be' items
                counter[3] += 1
                RC_textstringShould = RC_textstringShould + '\n' + paragraph.text
            else:
                if 'should not be' in paragraph.text:
                    # counting RC -> 'should not be' items
                    counter[4] += 1
                    RC_textstringShouldNot = RC_textstringShouldNot + '\n' + paragraph.text
                else:
                    # counting RC -> 'should not be' items
                    counter[5] += 1
                    RC_textstringOther = RC_textstringOther + '\n' + paragraph.text    

    #store the results into a TxT file
    with open(para_outputs[:-5]+'.txt', 'a+') as fw:
        fw.write(RL_textstringMust)
        fw.write(RL_textstringShall)
        fw.write(RL_textstringOther)
        fw.write(RC_textstringShould)
        fw.write(RC_textstringShouldNot)
        fw.write(RC_textstringOther)
        fw.write('\nRL-shall counting:' + str(counter[0]) + '\n')
        fw.write('RL-must not counting:' + str(counter[1]) + '\n')
        fw.write('RL-other counting:' + str(counter[2]) + '\n')
        fw.write('RC-should counting:' + str(counter[3]) + '\n')
        fw.write('RC-should not counting:' + str(counter[4]) + '\n')
        fw.write('RC-other counting:' + str(counter[5]) + '\n')
        fw.close()
    

#Interpret PDF file and store into Text file
def parsePDF(para_PDFfile_loc, para_results):
    
    import pdfminer
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfparser import PDFParser
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.converter import PDFPageAggregator
    from pdfminer.layout import LTTextBoxHorizontal, LAParams
    from pdfminer.pdfpage import PDFTextExtractionNotAllowed, PDFPage

    fp = open(para_PDFfile_loc, 'rb')

    # PDF parser instance
    parser = PDFParser(fp)
    doc = PDFDocument(parser)
 
    # not judget whether PDF is extractable or not since doc.is_extractable is justing checking that special tag
    # instead of checking the actual extractable status
    # if not doc.is_extractable:
    #    raise PDFTextExtractionNotAllowed
   
    # PDF resource manager
    resmgr = PDFResourceManager()

    # PDF Device
    laparams = LAParams()
    device = PDFPageAggregator(resmgr, laparams=laparams)

    # PDF page interpreter
    interpreter = PDFPageInterpreter(resmgr, device)
 
    #create an empty word document structure
    word_doc = Document()

    # go throu all pages
    # doc.get_pages() to pages
    for page in PDFPage.create_pages(doc):
        interpreter.process_page(page)

        layout = device.get_result()
        # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象
        # 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等
        # 想要获取文本就获得对象的text属性，
        for x in layout:
            if (isinstance(x, LTTextBoxHorizontal)):
                TXTinfo = x.get_text()
                word_doc.add_paragraph(TXTinfo)
    word_doc.save(para_results)

# GUI input to get the file to be interpreted
def PDFfile_input(para_promptinfo):

    # Define termination function 
    def return_callback(event):
        topwindow.quit()

    # Define close function for no input
    def close_callback():
        tkinter.messagebox.showinfo('message', 'No Input')
        # Terminate process
        topwindow.destroy()
        sys.exit()

    topwindow = tkinter.Tk()
    topwindow.title("PDF File Location")

    input_label = tkinter.Label(topwindow, text = para_promptinfo)
    input_label.pack(side = LEFT)

    input_entry = tkinter.Entry(topwindow, bd=5)
    input_entry.pack(side = RIGHT)

    # Bind "Return" keyinput with Entry box
    input_entry.bind('<Return>', return_callback)

    # Define window close event
    topwindow.protocol("WM_DELETE_WINDOW", close_callback)

    topwindow.mainloop()

    # Get window input on URL to be grapped
    PDF_file = tkinter.Entry.get(input_entry)

    #topwindow.destroy()
    #hide the input window
    #used specially for while-loop in case wrong value input
    topwindow.withdraw()

    return PDF_file

if __name__ == "__main__":
    main()